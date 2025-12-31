import re
from docx import Document
import os

def extract_entities(text):
    entities = {}

    # Aadhaar number pattern (with or without spaces)
    aadhaar = re.findall(r"\b\d{4}\s\d{4}\s\d{4}\b", text) or re.findall(r"\b\d{12}\b", text)
    if aadhaar:
        entities['Aadhaar'] = aadhaar[0]

    # PAN card pattern
    pan = re.findall(r"[A-Z]{5}[0-9]{4}[A-Z]{1}", text)
    if pan:
        entities['PAN'] = pan[0]

    # Date of Birth pattern (dd/mm/yyyy or dd-mm-yyyy)
    dob = re.findall(r"\b\d{2}[/-]\d{2}[/-]\d{4}\b", text)
    if dob:
        entities['DOB'] = dob[0]

    # Name (simple heuristic: first non-empty line with two or more words)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        if len(line.split()) >= 2 and not any(char.isdigit() for char in line):
            entities['Name'] = line
            break

    # Address detection (heuristic)
    address_lines = [line for line in lines if any(word in line for word in ['Road', 'Street', 'Nagar', 'Colony', 'City', 'District', 'Pin', 'PIN', 'Pincode'])]
    if address_lines:
        entities['Address'] = ', '.join(address_lines[:3])

    return entities

def fill_form(entities, output_folder):
    doc = Document()
    doc.add_heading('Government Service Form (Auto-Filled)', 0)

    if not entities:
        doc.add_paragraph('No entities were detected in the uploaded document.')
    else:
        for key, value in entities.items():
            doc.add_paragraph(f"{key}: {value}")

    os.makedirs(output_folder, exist_ok=True)
    filepath = os.path.join(output_folder, "filled_form.docx")
    doc.save(filepath)
    return filepath
